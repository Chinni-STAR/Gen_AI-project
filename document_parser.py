"""Extracts raw text from uploaded resume files (PDF, DOCX, TXT) with robust error handling."""

import io
import re
from pypdf import PdfReader
from docx import Document


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Extracts text from PDF, DOCX, or TXT file bytes.
    
    Raises ValueError for unsupported extensions, empty files, or scanned/image-only PDFs.
    """
    if not file_bytes or len(file_bytes) == 0:
        raise ValueError(f"Uploaded file '{filename}' is empty (0 bytes). Please upload a valid resume.")

    lower = filename.lower()

    if lower.endswith(".pdf"):
        text = _extract_pdf(file_bytes)
    elif lower.endswith(".docx"):
        text = _extract_docx(file_bytes)
    elif lower.endswith(".txt"):
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1", errors="ignore")
    else:
        raise ValueError(
            f"Unsupported file format: '{filename}'. Please upload a PDF (.pdf), Word document (.docx), or Text file (.txt)."
        )

    # Clean whitespace and control characters
    text = re.sub(r"[\r\t\x00-\x08\x0b\x0c\x0e-\x1f]", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if len(text) < 30:
        raise ValueError(
            f"Unable to extract readable text from '{filename}'. "
            "If this is a scanned PDF or image-only document, please convert it to a text-searchable PDF, DOCX, or TXT file."
        )

    return text


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                raise ValueError("PDF file is password-protected. Please upload an unencrypted document.")
        
        pages = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(page_text)

        return "\n\n".join(pages).strip()
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to read PDF file structure: {str(e)}")


def _extract_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    paragraphs.append(" | ".join(row_cells))

        return "\n\n".join(paragraphs).strip()
    except Exception as e:
        raise ValueError(f"Failed to read Word (.docx) document structure: {str(e)}")
